"""
Test inline math conversion in detail
"""
import sys
sys.path.insert(0, '.')

from docx import Document
from lxml import etree
from models.converter import convert_markdown_content_to_word
from utils import log

# Simple test with inline math
content = "Test formula: $f(x) = x^2$ and another $y = 2x$."

log.info("Converting simple inline math...")
output_path = 'outputs/test_inline_math.docx'

try:
    convert_markdown_content_to_word(content, output_path)
    log.info(f"✓ Created: {output_path}")
    
    # Analyze the output
    doc = Document(output_path)
    log.info(f"\nParagraphs: {len(doc.paragraphs)}")
    
    for i, para in enumerate(doc.paragraphs):
        log.info(f"\nParagraph {i}: {para.text}")
        
        # Check for math elements
        math_elems = para._element.findall('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath')
        log.info(f"  Math elements: {len(math_elems)}")
        
        if math_elems:
            for j, math_elem in enumerate(math_elems):
                xml_str = etree.tostring(math_elem, encoding='unicode')
                log.info(f"  Math {j}: {xml_str[:200]}...")
        
        # Check paragraph XML structure
        para_xml = etree.tostring(para._element, encoding='unicode')
        if 'oMath' in para_xml:
            log.info("  ✓ Contains oMath in XML")
        else:
            log.info("  ✗ No oMath in XML")
            
except Exception as e:
    log.error(f"✗ Error: {e}")
    import traceback
    traceback.print_exc()
