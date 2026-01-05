"""
Compare the generated DOCX with the reference DOCX
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from lxml import etree

print("="*80)
print("COMPARING GENERATED VS REFERENCE DOCX")
print("="*80)

# Load both documents
ref_doc = Document('latex-sample.docx')
new_doc = Document('outputs/latex-sample-test-new.docx')

print(f"\nReference DOCX paragraphs: {len(ref_doc.paragraphs)}")
print(f"Generated DOCX paragraphs: {len(new_doc.paragraphs)}")

# Compare paragraph 1 (has inline math)
print("\n" + "="*80)
print("PARAGRAPH 1 COMPARISON")
print("="*80)

ref_para = ref_doc.paragraphs[1]
new_para = new_doc.paragraphs[1]

print(f"\nReference text: {ref_para.text[:100]}")
print(f"Generated text: {new_para.text[:100]}")

# Find math elements
ref_math = ref_para._element.findall('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath')
new_math = new_para._element.findall('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath')

print(f"\nReference math elements: {len(ref_math)}")
print(f"Generated math elements: {len(new_math)}")

if new_math:
    print("\n" + "-"*80)
    print("FIRST GENERATED MATH ELEMENT (f(x)):")
    print("-"*80)
    xml_str = etree.tostring(new_math[0], encoding='unicode', pretty_print=True)
    print(xml_str[:800])
    
    if len(new_math) > 1:
        print("\n" + "-"*80)
        print("SECOND GENERATED MATH ELEMENT (piecewise):")
        print("-"*80)
        xml_str2 = etree.tostring(new_math[1], encoding='unicode', pretty_print=True)
        print(xml_str2[:1200])

print("\n" + "="*80)
print("VERIFICATION COMPLETE")
print("="*80)
print("\n✓ Generated DOCX file: outputs/latex-sample-test-new.docx")
print("✓ Please open both files in Microsoft Word to compare visually")
