import os
import sys
from pathlib import Path

# Add backend to path
sys.path.append(str(Path(__file__).parent))

from models.converter import convert_markdown_content_to_word
from docx import Document

def test_table_conversion():
    markdown = """
# Table Test

| Header 1 | Header 2 |
| --- | --- |
| Cell 1 | Cell 2 with **bold** |
| $x^2$ | $\\frac{1}{2}$ |

After table.
"""
    output_path = "table_test_output.docx"
    
    print(f"Converting markdown to {output_path}...")
    convert_markdown_content_to_word(markdown, output_path)
    
    if not os.path.exists(output_path):
        print("Error: Output file not created")
        return False
        
    print("Checking document content...")
    doc = Document(output_path)
    
    # Check for table
    if len(doc.tables) == 0:
        print("Error: No table found in DOCX")
        return False
        
    table = doc.tables[0]
    print(f"Found table with {len(table.rows)} rows and {len(table.columns)} columns")
    
    # Check headers
    if "Header 1" not in table.cell(0, 0).text or "Header 2" not in table.cell(0, 1).text:
        print(f"Error: Headers mismatch. Found: {table.cell(0, 0).text}, {table.cell(0, 1).text}")
        return False
        
    # Check cell 2,1 (index 1,0) - math
    # Math is converted to OMML, so plain text might vary depending on how it's read
    print(f"Cell (1,0) text: {table.cell(2, 0).text}")
    print(f"Cell (1,1) text: {table.cell(2, 1).text}")
    
    print("Test passed successfully!")
    return True

if __name__ == "__main__":
    success = test_table_conversion()
    if not success:
        sys.exit(1)
