"""
Markdown to DOCX converter with LaTeX formula support
"""
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional
import re
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from markdown_it import MarkdownIt
from mdit_py_plugins.texmath import texmath_plugin
from latex2mathml.converter import convert as latex_to_mathml
from lxml import etree

from utils import log


class ListManager:
    """Manages list numbering for Word documents"""
    
    def __init__(self, doc: Document):
        self.doc = doc
        self.list_count = 0
        self._ensure_numbering()
    
    def _ensure_numbering(self):
        """Ensure the document has numbering part and basic abstract definitions"""
        try:
            # This triggers the creation of numbering.xml if it doesn't exist
            # and defines the standard styles like 'List Number'
            self.doc.part.numbering_part
        except Exception:
            pass

    def get_new_num_id(self, is_ordered: bool = True) -> int:
        """Create a new numbering definition and return its numId"""
        self.list_count += 1
        
        # Access the underlying XML
        numbering = self.doc.part.numbering_part.numbering_definitions._numbering
        
        # 1. Create a new abstractNum
        abstract_num_id = self.list_count
        abstractNum = OxmlElement('w:abstractNum')
        abstractNum.set(qn('w:abstractNumId'), str(abstract_num_id))
        
        # Add basic multi-level support (Word expects 9 levels)
        for level in range(9):
            lvl = OxmlElement('w:lvl')
            lvl.set(qn('w:ilvl'), str(level))
            
            start = OxmlElement('w:start')
            start.set(qn('w:val'), '1')
            lvl.append(start)
            
            numFmt = OxmlElement('w:numFmt')
            if is_ordered:
                numFmt.set(qn('w:val'), 'decimal')
            else:
                numFmt.set(qn('w:val'), 'bullet')
            lvl.append(numFmt)
            
            lvlText = OxmlElement('w:lvlText')
            if is_ordered:
                # Use standard %1. format for level 0
                text = f"%{level + 1}."
                lvlText.set(qn('w:val'), text)
            else:
                lvlText.set(qn('w:val'), '•')
            lvl.append(lvlText)
            
            lvlJc = OxmlElement('w:lvlJc')
            lvlJc.set(qn('w:val'), 'left')
            lvl.append(lvlJc)
            
            # Indentation
            pPr = OxmlElement('w:pPr')
            ind = OxmlElement('w:ind')
            # 360 twips (0.25 inch) per level increment
            left = 720 + (level * 360) 
            hanging = 360
            ind.set(qn('w:left'), str(left))
            ind.set(qn('w:hanging'), str(hanging))
            pPr.append(ind)
            lvl.append(pPr)
            
            abstractNum.append(lvl)
        
        # Insert abstractNum into numbering.xml (must be before 'num' elements)
        # Find the first 'num' element or end of list
        nums = numbering.xpath('w:num')
        if nums:
            nums[0].addprevious(abstractNum)
        else:
            numbering.append(abstractNum)
            
        # 2. Create a new num (instance) that points to the abstractNum
        num = OxmlElement('w:num')
        num.set(qn('w:numId'), str(self.list_count + 100)) # Offset to avoid conflict
        
        abstractNumId = OxmlElement('w:abstractNumId')
        abstractNumId.set(qn('w:val'), str(abstract_num_id))
        num.append(abstractNumId)
        
        numbering.append(num)
        
        return self.list_count + 100

    @staticmethod
    def set_paragraph_numbering(paragraph, num_id: int, level: int = 0):
        """Apply numbering attributes to a paragraph's XML"""
        pPr = paragraph._element.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), str(level))
        numPr.append(ilvl)
        
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), str(num_id))
        numPr.append(numId)
        
        pPr.append(numPr)


class AnyXmlElement:
    """Helper class to create arbitrary XML elements for OMML"""
    
    def __init__(self, tag: str, text: Optional[str] = None, **attrs):
        self.tag = tag
        self.text = text
        self.attrs = attrs
        self.children = []
    
    def add_child(self, child):
        self.children.append(child)
        return self
    
    def to_element(self) -> OxmlElement:
        """Convert to docx OxmlElement"""
        elem = OxmlElement(self.tag)
        
        for key, value in self.attrs.items():
            elem.set(qn(key), str(value))
        
        if self.text:
            elem.text = self.text
        
        for child in self.children:
            if isinstance(child, AnyXmlElement):
                elem.append(child.to_element())
            elif isinstance(child, OxmlElement):
                elem.append(child)
        
        return elem


def mathml_to_omml(mathml_str: str) -> Optional[OxmlElement]:
    """
    Convert MathML to Office Math Markup Language (OMML)
    Enhanced version with comprehensive MathML element support
    """
    if not mathml_str:
        return None
        
    try:
        # Parse MathML
        try:
            # Clean up the mathml string if it has encoding declarations that might conflict
            if isinstance(mathml_str, str):
                mathml_bytes = mathml_str.encode('utf-8')
            else:
                mathml_bytes = mathml_str
            
            mathml = etree.fromstring(mathml_bytes)
        except etree.XMLSyntaxError as e:
            log.error(f"XML Syntax Error in MathML: {e}")
            # Try a more lenient parse if possible or just fail gracefully
            return None
        
        # Create OMML root element with proper namespace
        omml = OxmlElement('m:oMath')
        # Note: Don't set xmlns:m here as it's already defined in the qn() calls
        
        def create_run(text: str, style: str = 'p') -> OxmlElement:
            """Create an OMML run element with text and style
            
            Args:
                text: The text content
                style: 'i' for italic (variables), 'p' for plain (operators, numbers)
            """
            r = OxmlElement('m:r')
            
            # Add run properties with style
            rPr = OxmlElement('m:rPr')
            sty = OxmlElement('m:sty')
            sty.set(qn('m:val'), style)
            rPr.append(sty)
            r.append(rPr)
            
            # Add text
            t = OxmlElement('m:t')
            t.text = text or ''
            # Preserve spaces
            if text and (text.startswith(' ') or text.endswith(' ')):
                t.set(qn('xml:space'), 'preserve')
            r.append(t)
            
            return r
        
        def convert_element(elem, parent_omml):
            """Recursively convert MathML elements to OMML"""
            tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
            
            if tag == 'mn':
                # Number - plain style
                parent_omml.append(create_run(elem.text or '', 'p'))
            
            elif tag == 'mi':
                # Identifier (variable) - italic style
                parent_omml.append(create_run(elem.text or '', 'i'))
            
            elif tag == 'mo':
                # Operator - plain style
                parent_omml.append(create_run(elem.text or '', 'p'))
            
            elif tag == 'mtext':
                # Text in math - plain style
                parent_omml.append(create_run(elem.text or '', 'p'))
            
            elif tag == 'mspace':
                # Space - plain style
                parent_omml.append(create_run(' ', 'p'))
            
            elif tag == 'mfrac':
                # Fraction
                frac = OxmlElement('m:f')
                num = OxmlElement('m:num')
                den = OxmlElement('m:den')
                
                children = list(elem)
                if len(children) >= 2:
                    convert_element(children[0], num)
                    convert_element(children[1], den)
                
                frac.append(num)
                frac.append(den)
                parent_omml.append(frac)
            
            elif tag == 'msup':
                # Superscript
                sup = OxmlElement('m:sSup')
                base = OxmlElement('m:e')
                supElem = OxmlElement('m:sup')
                
                children = list(elem)
                if len(children) >= 2:
                    convert_element(children[0], base)
                    convert_element(children[1], supElem)
                
                sup.append(base)
                sup.append(supElem)
                parent_omml.append(sup)
            
            elif tag == 'msub':
                # Subscript
                sub = OxmlElement('m:sSub')
                base = OxmlElement('m:e')
                subElem = OxmlElement('m:sub')
                
                children = list(elem)
                if len(children) >= 2:
                    convert_element(children[0], base)
                    convert_element(children[1], subElem)
                
                sub.append(base)
                sub.append(subElem)
                parent_omml.append(sub)
            
            elif tag == 'msubsup':
                # Subscript and superscript
                sSubSup = OxmlElement('m:sSubSup')
                base = OxmlElement('m:e')
                subElem = OxmlElement('m:sub')
                supElem = OxmlElement('m:sup')
                
                children = list(elem)
                if len(children) >= 3:
                    convert_element(children[0], base)
                    convert_element(children[1], subElem)
                    convert_element(children[2], supElem)
                
                sSubSup.append(base)
                sSubSup.append(subElem)
                sSubSup.append(supElem)
                parent_omml.append(sSubSup)
            
            elif tag == 'msqrt':
                # Square root
                rad = OxmlElement('m:rad')
                radPr = OxmlElement('m:radPr')
                degHide = OxmlElement('m:degHide')
                degHide.set(qn('m:val'), '1')
                radPr.append(degHide)
                
                base = OxmlElement('m:e')
                for child in elem:
                    convert_element(child, base)
                
                rad.append(radPr)
                rad.append(base)
                parent_omml.append(rad)
            
            elif tag == 'mroot':
                # Nth root
                rad = OxmlElement('m:rad')
                deg = OxmlElement('m:deg')
                base = OxmlElement('m:e')
                
                children = list(elem)
                if len(children) >= 2:
                    convert_element(children[0], base)
                    convert_element(children[1], deg)
                
                rad.append(deg)
                rad.append(base)
                parent_omml.append(rad)
            
            elif tag == 'munder':
                # Underscript
                func = OxmlElement('m:func')
                fName = OxmlElement('m:fName')
                base = OxmlElement('m:e')
                
                children = list(elem)
                if len(children) >= 2:
                    convert_element(children[0], fName)
                    convert_element(children[1], base)
                
                func.append(fName)
                func.append(base)
                parent_omml.append(func)
            
            elif tag == 'mover':
                # Overscript (like accent)
                acc = OxmlElement('m:acc')
                accPr = OxmlElement('m:accPr')
                base = OxmlElement('m:e')
                
                children = list(elem)
                if len(children) >= 2:
                    convert_element(children[0], base)
                    # Second child is the accent character
                    if children[1].text:
                        chrElem = OxmlElement('m:chr')
                        chrElem.set(qn('m:val'), children[1].text)
                        accPr.append(chrElem)
                
                acc.append(accPr)
                acc.append(base)
                parent_omml.append(acc)
            
            elif tag == 'munderover':
                # Under and over script (like limits)
                limLow = OxmlElement('m:limLow')
                limUpp = OxmlElement('m:limUpp')
                base = OxmlElement('m:e')
                lim = OxmlElement('m:lim')
                
                children = list(elem)
                if len(children) >= 3:
                    convert_element(children[0], base)
                    # Create nested structure for under and over
                    innerBase = OxmlElement('m:e')
                    convert_element(children[0], innerBase)
                    limLow.append(innerBase)
                    
                    underLim = OxmlElement('m:lim')
                    convert_element(children[1], underLim)
                    limLow.append(underLim)
                    
                    limUpp.append(limLow)
                    overLim = OxmlElement('m:lim')
                    convert_element(children[2], overLim)
                    limUpp.append(overLim)
                    
                    parent_omml.append(limUpp)
                else:
                    # Fallback
                    for child in children:
                        convert_element(child, parent_omml)
            
            elif tag == 'mfenced':
                # Fenced expression (parentheses, brackets, etc.)
                d = OxmlElement('m:d')
                dPr = OxmlElement('m:dPr')
                
                # Get opening and closing characters
                open_char = elem.get('open', '(')
                close_char = elem.get('close', ')')
                
                begChr = OxmlElement('m:begChr')
                begChr.set(qn('m:val'), open_char)
                dPr.append(begChr)
                
                endChr = OxmlElement('m:endChr')
                endChr.set(qn('m:val'), close_char)
                dPr.append(endChr)
                
                d.append(dPr)
                
                # Add content
                base = OxmlElement('m:e')
                for child in elem:
                    convert_element(child, base)
                d.append(base)
                
                parent_omml.append(d)
            
            elif tag == 'mtable':
                # Matrix/Table
                matrix = OxmlElement('m:m')
                
                for row_elem in elem:
                    if row_elem.tag.split('}')[-1] == 'mtr':
                        mr = OxmlElement('m:mr')
                        for cell_elem in row_elem:
                            if cell_elem.tag.split('}')[-1] == 'mtd':
                                e = OxmlElement('m:e')
                                for child in cell_elem:
                                    convert_element(child, e)
                                mr.append(e)
                        matrix.append(mr)
                
                parent_omml.append(matrix)
            
            elif tag == 'mrow' or tag == 'math':
                # Row or root - process children
                for child in elem:
                    convert_element(child, parent_omml)
            
            else:
                # Default: process children
                for child in elem:
                    convert_element(child, parent_omml)
        
        # Convert the MathML tree
        convert_element(mathml, omml)
        
        return omml
    
    except Exception as e:
        log.error(f"Error converting MathML to OMML: {e}")
        return None


def convert_latex_to_omml(latex: str, is_block: bool = False) -> Optional[OxmlElement]:
    """Convert LaTeX formula to OMML for Word"""
    try:
        if not latex:
            return None
        
        # Convert LaTeX to MathML
        try:
            mathml = latex_to_mathml(latex)
        except Exception as e:
            log.error(f"latex2mathml conversion failed for: {latex[:50]}... Error: {e}")
            return None
        
        # Convert MathML to OMML
        try:
            omml = mathml_to_omml(mathml)
        except Exception as e:
            log.error(f"mathml_to_omml failed. Error: {e}")
            return None
        
        return omml
    
    except Exception as e:
        log.error(f"General error in convert_latex_to_omml: {e}")
        return None


def parse_markdown(content: str) -> List[Dict[str, Any]]:
    """Parse markdown content to tokens"""
    log.info("Parsing Markdown content")
    
    try:
        # Initialize markdown-it with LaTeX support
        md = (
            MarkdownIt('commonmark', {'breaks': True, 'html': True})
            .use(texmath_plugin, delimiters='dollars')
            .enable('table')
        )
        
        # Parse to tokens
        tokens = md.parse(content)
        
        log.info(f"Parsed {len(tokens)} tokens from Markdown")
        return tokens
    
    except Exception as e:
        log.error(f"Error parsing Markdown: {e}")
        raise Exception(f"Failed to parse Markdown: {e}")


def tokens_to_docx_paragraphs(doc: Document, tokens: List[Dict[str, Any]]) -> Tuple[List, List]:
    """Convert markdown tokens to Word document paragraphs"""
    paragraphs = []
    numbering_configs = []
    list_level = 0
    list_stack = []
    
    # Initialize numbering manager
    list_manager = ListManager(doc)
    
    i = 0
    while i < len(tokens):
        token = tokens[i]
        token_type = token.type
        
        if token_type == 'heading_open':
            level = int(token.tag[1])  # h1 -> 1, h2 -> 2, etc.
            
            # Get the inline content
            if i + 1 < len(tokens) and tokens[i + 1].type == 'inline':
                heading_text = tokens[i + 1].content
                
                para = doc.add_paragraph(heading_text)
                para.style = f'Heading {level}'
                paragraphs.append(para)
            
            i += 2  # Skip inline and heading_close
            continue
        
        elif token_type == 'paragraph_open':
            if i + 1 < len(tokens) and tokens[i + 1].type == 'inline':
                inline_token = tokens[i + 1]
                
                para = doc.add_paragraph()
                parse_inline_content(para, inline_token)
                paragraphs.append(para)
            
            i += 2  # Skip inline and paragraph_close
            continue
        
        elif token_type == 'bullet_list_open' or token_type == 'ordered_list_open':
            list_level += 1
            is_ordered = token_type == 'ordered_list_open'
            
            # Create a new numbering instance for this list
            num_id = list_manager.get_new_num_id(is_ordered)
            
            list_stack.append({
                'type': 'ordered' if is_ordered else 'bullet',
                'level': list_level - 1,
                'num_id': num_id
            })
        
        elif token_type == 'bullet_list_close' or token_type == 'ordered_list_close':
            list_level -= 1
            if list_stack:
                list_stack.pop()
        
        elif token_type == 'list_item_open':
            # Find the list info from stack
            list_info = list_stack[-1] if list_stack else None
            
            # Process list item content
            i += 1
            # Skip potential paragraph_open within list item to keep it simple
            if i < len(tokens) and tokens[i].type == 'paragraph_open':
                i += 1
                if i < len(tokens) and tokens[i].type == 'inline':
                    inline_token = tokens[i]
                    
                    para = doc.add_paragraph()
                    parse_inline_content(para, inline_token)
                    
                    if list_info:
                        # Apply style for basic formatting
                        if list_info['type'] == 'ordered':
                            para.style = 'List Number'
                        else:
                            para.style = 'List Bullet'
                            
                        # Apply unique numbering to force reset and level
                        ListManager.set_paragraph_numbering(
                            para, 
                            list_info['num_id'], 
                            list_info['level']
                        )
                    
                    paragraphs.append(para)
            
            # Continue will handle the skipping of tokens within the list item
            continue
        
        elif token_type == 'fence' or token_type == 'code_block':
            para = doc.add_paragraph(token.content)
            para.style = 'No Spacing'
            
            for run in para.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            
            paragraphs.append(para)
        
        elif token_type == 'hr':
            para = doc.add_paragraph('─' * 50)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraphs.append(para)
        
        elif token_type == 'math_block' or token_type == 'math_block_end':
            if hasattr(token, 'content') and token.content:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Try to add OMML math
                omml = convert_latex_to_omml(token.content, is_block=True)
                if omml:
                    para._element.append(omml)
                else:
                    # Fallback to plain text
                    para.add_run(token.content)
                
                paragraphs.append(para)
        
        i += 1
    
    log.info(f"Converted tokens to {len(paragraphs)} paragraphs")
    return paragraphs, numbering_configs


def parse_inline_content(paragraph, inline_token) -> None:
    """Parse inline content and add runs to paragraph"""
    if not hasattr(inline_token, 'children') or not inline_token.children:
        if hasattr(inline_token, 'content') and inline_token.content:
            paragraph.add_run(inline_token.content)
        return
    
    for child in inline_token.children:
        child_type = child.type
        
        if child_type == 'text':
            paragraph.add_run(child.content)
        
        elif child_type == 'strong':
            run = paragraph.add_run(child.content)
            run.bold = True
        
        elif child_type == 'em':
            run = paragraph.add_run(child.content)
            run.italic = True
        
        elif child_type == 'code_inline':
            run = paragraph.add_run(child.content)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        
        elif child_type == 'softbreak' or child_type == 'hardbreak':
            paragraph.add_run('\n')
        
        elif child_type == 'math_inline':
            # Try to add inline math
            omml = convert_latex_to_omml(child.content, is_block=False)
            if omml:
                paragraph._element.append(omml)
            else:
                # Fallback to plain text
                paragraph.add_run(f"${child.content}$")
        
        else:
            # Default: add as text if has content
            if hasattr(child, 'content') and child.content:
                paragraph.add_run(child.content)


def convert_markdown_to_word(input_path: str, output_path: str) -> str:
    """Convert Markdown file to Word document"""
    log.info(f"Converting Markdown to Word: {input_path} -> {output_path}")
    
    try:
        # Read markdown file
        with open(input_path, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        log.info(f"Read {len(markdown_content)} characters from input file")
        
        # Parse markdown
        tokens = parse_markdown(markdown_content)
        
        # Create Word document
        doc = Document()
        
        # Convert tokens to paragraphs
        paragraphs, numbering_configs = tokens_to_docx_paragraphs(doc, tokens)
        
        # Save document
        doc.save(output_path)
        log.info(f"Successfully created Word document: {output_path}")
        
        return output_path
    
    except Exception as e:
        log.error(f"Error converting Markdown to Word: {e}")
        raise Exception(f"Conversion failed: {e}")


def convert_markdown_content_to_word(content: str, output_path: str) -> str:
    """Convert Markdown content to Word document"""
    log.info(f"Converting Markdown content to Word. Output: {output_path}")
    
    try:
        # Parse markdown
        log.debug("Step 1: Parsing tokens")
        tokens = parse_markdown(content)
        
        # Create Word document
        log.debug("Step 2: Initializing Document")
        doc = Document()
        
        # Convert tokens to paragraphs
        log.debug(f"Step 3: Converting {len(tokens)} tokens to paragraphs")
        paragraphs, numbering_configs = tokens_to_docx_paragraphs(doc, tokens)
        
        # Save document
        log.debug("Step 4: Saving document")
        doc.save(output_path)
        log.info(f"Successfully created Word document from content: {output_path}")
        
        return output_path
    
    except Exception as e:
        log.error(f"Error in convert_markdown_content_to_word: {str(e)}", exc_info=True)
        # Ensure we return a clean string for the exception to avoid serialization issues
        raise Exception(f"Conversion failed at internal step. Technical details: {type(e).__name__}")
